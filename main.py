import re
import random
from docx import Document
from docx.shared import Pt

def load_passages_from_txt(filename="passages.txt"):
    with open(filename, 'r', encoding='utf-8') as f:
        text = f.read()
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    passages = re.split(r'\n{3,}', text)
    return [p.strip() for p in passages if p.strip()]

def split_sentences(passage):
    # 강화된 문장 분리 기준: 마침표, 느낌표, 물음표 다음 공백
    pattern = r'(?<=[.!?])(?<!\.\.\.)(?<![A-Z][a-z]\.)(?<!\b[A-Z]\.)(?<!\.\w)(?=\s)'
    sentences = re.split(pattern, passage)
    return [s.strip() for s in sentences if s.strip()]

def num_to_circle(num):
    if 1 <= num <= 20:
        return chr(0x2460 + num - 1)
    else:
        return f"({num})"

def create_questions_and_answers(passages):
    questions_doc = Document()
    answers_doc = Document()

    answers_doc.add_paragraph("※ 답안지\n")

    for idx, passage in enumerate(passages, 1):
        sentences = split_sentences(passage)

        if len(sentences) < 2:
            continue  # 문장이 너무 짧으면 건너뜀

        # 무작위 위치 선정 (첫 문장은 제외)
        missing_index = random.randint(1, len(sentences) - 1)
        missing_sentence = sentences.pop(missing_index)

        # 문제 본문 생성
        marked_sentences = []
        for i, sent in enumerate(sentences, 1):
            punctuation = '' if sent.endswith(('.', '?', '!')) else '.'
            marked_sentences.append(f"{sent}{punctuation} {num_to_circle(i)}")

        problem_text = ' '.join(marked_sentences)

        # 문제 작성
        questions_doc.add_paragraph(f"{idx}. 다음 문장이 들어갈 곳으로 알맞은 곳은?")
        questions_doc.add_paragraph("")  # 빈 줄
        questions_doc.add_paragraph(missing_sentence)
        questions_doc.add_paragraph("")  # 빈 줄
        questions_doc.add_paragraph(problem_text)
        questions_doc.add_paragraph("")  # 문제 간 간격

        # 답안 작성
        answers_doc.add_paragraph(f"{idx}. {num_to_circle(missing_index)}")

    questions_doc.save("questions.docx")
    answers_doc.save("answers.docx")

if __name__ == "__main__":
    passages = load_passages_from_txt("passages.txt")
    create_questions_and_answers(passages)
